#1

import streamlit as st
from PyPDF2 import PdfReader
import docx
import os
import base64
import random
import string
from io import BytesIO
from gtts import gTTS
import matplotlib.pyplot as plt

# For Email (SMTP placeholder)
import smtplib
from email.mime.text import MIMEText

# For AI Feedback (Groq API)
import requests

# Session state initialization
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'user_email' not in st.session_state:
    st.session_state.user_email = ""
if 'is_admin' not in st.session_state:
    st.session_state.is_admin = False

# Set page config
st.set_page_config(page_title="AI Resume Feedback Bot", page_icon="📄", layout="wide")

# ------------------- Helper Functions -------------------

# Read content from uploaded resume (PDF/DOCX)
def extract_text_from_resume(uploaded_file):
    if uploaded_file.type == "application/pdf":
        pdf = PdfReader(uploaded_file)
        return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        return ""

# Simulated OTP Sender
def send_otp(email_or_phone):
    otp = ''.join(random.choices(string.digits, k=6))
    st.session_state.otp_sent = otp
    # Placeholder - log OTP
    print(f"📨 OTP sent to {email_or_phone}: {otp}")
    return otp

# Placeholder: Send Email (SMTP)
def send_email(receiver_email, subject, body):
    print(f"Sending email to {receiver_email}...\nSubject: {subject}\n{body}")
    # SMTP logic placeholder here
    return True

# Convert audio from text
def generate_audio(text):
    tts = gTTS(text)
    audio_fp = BytesIO()
    tts.write_to_fp(audio_fp)
    audio_fp.seek(0)
    return audio_fp

# Base64 download button (PDF export / audio)
def get_download_button(file_data, filename, label):
    b64 = base64.b64encode(file_data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{label}</a>'
    return href


#2
# ------------------- UI Styling -------------------
st.markdown("""
    <style>
        .main {background-color: #f5f5f5;}
        h1, h2, h3 {color: #333;}
        .stButton>button {
            background-color: #4CAF50; color: white; border: none;
            padding: 0.5em 1.5em; border-radius: 5px;
        }
        .stTextInput>div>div>input {
            border: 1px solid #ccc; padding: 0.4em;
        }
    </style>
""", unsafe_allow_html=True)

# ------------------- Login / Signup -------------------
def login_signup_screen():
    st.title("🔐 Login or Signup to Continue")
    tabs = st.tabs(["🔑 Login", "📝 Signup"])

    with tabs[0]:
        email = st.text_input("📧 Email", key="login_email")
        password = st.text_input("🔒 Password", type="password", key="login_password")

        if st.button("🔓 Login"):
            if os.path.exists("users.txt"):
                with open("users.txt", "r") as f:
                    for line in f.readlines():
                        saved_name, saved_email, saved_phone, saved_pass = line.strip().split("|")
                        if saved_email == email and saved_pass == password:
                            st.success(f"✅ Welcome back, {saved_name}!")
                            st.session_state.logged_in = True
                            st.session_state.user_email = email
                            st.session_state.is_admin = (email == "admin@gmail.com")
                            return
                st.error("❌ Incorrect credentials.")
            else:
                st.error("❌ No users found.")

    with tabs[1]:
        name = st.text_input("👤 Full Name")
        signup_email = st.text_input("📧 Email", key="signup_email")
        phone = st.text_input("📱 Phone")
        new_password = st.text_input("🔒 Create Password", type="password")

        if st.button("📩 Send OTP"):
            if signup_email or phone:
                send_otp(signup_email or phone)
                st.session_state.otp_step = True
            else:
                st.warning("⚠️ Enter email or phone to receive OTP.")

        if st.session_state.get("otp_step"):
            otp_input = st.text_input("🔑 Enter OTP")
            if st.button("✅ Verify and Signup"):
                if otp_input == st.session_state.otp_sent:
                    # Save user
                    with open("users.txt", "a") as f:
                        f.write(f"{name}|{signup_email}|{phone}|{new_password}\n")
                    st.success("🎉 Signup successful. You can now login.")
                    st.session_state.otp_step = False
                else:
                    st.error("❌ Invalid OTP.")

# Show login/signup screen only if not logged in
if not st.session_state.logged_in:
    login_signup_screen()
    st.stop()


#3
# ------------------- Dashboard Screen -------------------
def dashboard():
    st.markdown("<h2 style='color:#4CAF50'>📊 Welcome to Resume Feedback Bot</h2>", unsafe_allow_html=True)

    dashboard_menu = st.sidebar.selectbox(
        "📂 Dashboard Menu",
        [
            "📤 Upload Resume",
            "📈 Resume Feedback",
            "🔄 Rewritten Resume",
            "🔈 Audio Tips",
            "📧 Email Feedback",
            "📂 My Feedback History",
            "🔒 Change Password",
            "❓ Forgot Password",
            "👑 Admin Dashboard" if st.session_state.is_admin else "---"
        ]
    )

    if dashboard_menu == "📤 Upload Resume":
        upload_resume_section()
    elif dashboard_menu == "📈 Resume Feedback":
        resume_feedback_section()
    elif dashboard_menu == "🔄 Rewritten Resume":
        rewritten_resume_section()
    elif dashboard_menu == "🔈 Audio Tips":
        audio_tips_section()
    elif dashboard_menu == "📧 Email Feedback":
        email_feedback_section()
    elif dashboard_menu == "📂 My Feedback History":
        feedback_history_section()
    elif dashboard_menu == "🔒 Change Password":
        change_password_section()
    elif dashboard_menu == "❓ Forgot Password":
        forgot_password_section()
    elif dashboard_menu == "👑 Admin Dashboard":
        admin_dashboard_section()

# ------------------- Show Dashboard only after login -------------------
if st.session_state.logged_in:
    dashboard()


#4
# ------------------- Upload Resume Section -------------------
def upload_resume_section():
    st.markdown("## 📤 Upload Your Resume")
    st.markdown("Upload your resume in **PDF** or **Word (DOCX)** format.")
    
    uploaded_file = st.file_uploader("Choose a resume file", type=["pdf", "docx"])

    if uploaded_file:
        file_extension = uploaded_file.name.split('.')[-1]

        if file_extension == "pdf":
            resume_text = extract_text_from_pdf(uploaded_file)
        elif file_extension == "docx":
            resume_text = extract_text_from_docx(uploaded_file)
        else:
            st.error("Unsupported file format.")
            return

        st.session_state.resume_text = resume_text
        st.success("✅ Resume uploaded and text extracted!")

        st.markdown("### 🎯 Enter Your Target Job Role")
        st.session_state.target_role = st.text_input("e.g., Data Scientist, Software Engineer")

        if st.button("🔍 Proceed to Get Feedback"):
            if st.session_state.target_role.strip() == "":
                st.warning("Please enter a target job role.")
            else:
                st.session_state.move_to_feedback = True
                st.success("Ready to fetch feedback. Go to Resume Feedback section.")
    else:
        st.info("No file uploaded yet.")

# ------------------- Extract Text from PDF -------------------
def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text

# ------------------- Extract Text from DOCX -------------------
def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])


#5
# ------------------- Resume Feedback via Groq -------------------
def generate_resume_feedback():
    st.markdown("## 🤖 AI-Powered Resume Feedback")
    if not st.session_state.get("resume_text") or not st.session_state.get("target_role"):
        st.warning("Please upload a resume and enter your target role first.")
        return

    with st.spinner("Analyzing your resume with AI..."):
        prompt = f"""You are a professional resume reviewer.
The candidate is applying for the role of **{st.session_state.target_role}**.
Here is their resume content:

{st.session_state.resume_text}

Please provide:
1. A resume score out of 100.
2. Top 3 strengths.
3. Top 3 weaknesses.
4. One tip to improve the resume.
5. ATS (Applicant Tracking System) friendly suggestion."""

        try:
            response = groq.ChatCompletion.create(
                model="llama3-70b-8192",
                messages=[{"role": "user", "content": prompt}]
            )
            result = response.choices[0].message.content
            st.session_state.feedback = result
        except Exception as e:
            st.error(f"Groq API error: {e}")
            return

    display_feedback(st.session_state.feedback)

# ------------------- Display Feedback with Score -------------------
def display_feedback(feedback):
    st.success("✅ AI Feedback Generated!")
    score = extract_score(feedback)

    if score:
        st.markdown("### 📊 Resume Score")
        st.progress(score / 100)
        st.markdown(f"**Score: {score}/100**")

    st.markdown("### 📋 Full Feedback")
    st.text_area("📝 AI Review", feedback, height=300)

# ------------------- Extract Score from Feedback -------------------
def extract_score(feedback):
    match = re.search(r"score(?:\s+)?(?:is)?(?:\s+)?:?(\d{1,3})", feedback, re.IGNORECASE)
    if not match:
        match = re.search(r"(\d{1,3})\s*/\s*100", feedback)
    if match:
        score = int(match.group(1))
        return min(score, 100)
    return None


#6
# ------------------- AI Resume Rewriting -------------------
def rewrite_resume():
    st.markdown("## 🔄 Rewritten Resume (AI-Powered)")
    if not st.session_state.get("resume_text") or not st.session_state.get("target_role"):
        st.warning("Please upload a resume and enter target role.")
        return

    with st.spinner("Generating rewritten resume..."):
        prompt = f"""You are an expert resume writer.
Rephrase and rewrite this resume for the role of **{st.session_state.target_role}**.
Keep it professional and ATS-friendly.

Resume:
{st.session_state.resume_text}"""

        try:
            response = groq.ChatCompletion.create(
                model="llama3-70b-8192",
                messages=[{"role": "user", "content": prompt}]
            )
            rewritten = response.choices[0].message.content
            st.session_state.rewritten_resume = rewritten
        except Exception as e:
            st.error(f"Groq error: {e}")
            return

    st.success("✅ Resume Rewritten Successfully!")
    st.text_area("📝 Rewritten Resume", rewritten, height=400)
    st.download_button("📥 Download Rewritten Resume", data=rewritten, file_name="rewritten_resume.txt")

# ------------------- Voice Feedback (gTTS) -------------------
def generate_audio_feedback():
    st.markdown("## 🔈 Voice Tips")
    if not st.session_state.get("feedback"):
        st.warning("Generate resume feedback first.")
        return

    with st.spinner("Generating voice tips..."):
        try:
            tts = gTTS(st.session_state.feedback)
            audio_path = "voice_feedback.mp3"
            tts.save(audio_path)
            st.audio(audio_path)
            with open(audio_path, "rb") as f:
                st.download_button("📥 Download Voice Tips", f, file_name="voice_feedback.mp3")
        except Exception as e:
            st.error(f"Audio generation error: {e}")

# ------------------- Send Feedback via Email (SMTP Placeholder) -------------------
def send_email_feedback():
    st.markdown("## 📧 Email Your Feedback")
    email = st.text_input("Enter your email")
    if st.button("📨 Send Email"):
        if not email or "@" not in email:
            st.error("Enter a valid email.")
            return
        if not st.session_state.get("feedback"):
            st.warning("Generate feedback first.")
            return
        try:
            # SMTP Placeholder logic — use your SMTP credentials here
            msg = EmailMessage()
            msg["Subject"] = "Your AI Resume Feedback"
            msg["From"] = "noreply@example.com"
            msg["To"] = email
            msg.set_content(st.session_state.feedback)

            # Example: use smtp.gmail.com with app password
            server = smtplib.SMTP("smtp.gmail.com", 587)
            server.starttls()
            server.login("your_email@gmail.com", "your_app_password")  # Use .env in real deployment
            server.send_message(msg)
            server.quit()
            st.success("✅ Feedback sent successfully!")
        except Exception as e:
            st.error(f"Failed to send email: {e}")


#7
# ------------------- Feedback History -------------------
def view_feedback_history():
    st.markdown("## 📂 Feedback History")
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT id, timestamp, role, score FROM feedback WHERE email=?", (st.session_state.email,))
    rows = c.fetchall()
    conn.close()

    if rows:
        for row in rows:
            fid, ts, role, score = row
            st.markdown(f"📄 **Role:** {role} | 🕒 {ts} | 🧠 Score: {score}")
            col1, col2 = st.columns([4, 1])
            with col1:
                if st.button(f"🔍 View Feedback {fid}"):
                    show_feedback_by_id(fid)
            with col2:
                if st.button(f"🗑️ Delete {fid}"):
                    delete_feedback_by_id(fid)
    else:
        st.info("No previous feedback found.")

def show_feedback_by_id(fid):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT feedback, rewritten FROM feedback WHERE id=? AND email=?", (fid, st.session_state.email))
    result = c.fetchone()
    conn.close()
    if result:
        feedback_text, rewritten_text = result
        st.markdown("### 🧠 Feedback")
        st.text_area("Feedback", feedback_text, height=200)
        st.markdown("### ✍️ Rewritten Resume")
        st.text_area("Rewritten", rewritten_text, height=300)

def delete_feedback_by_id(fid):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM feedback WHERE id=? AND email=?", (fid, st.session_state.email))
    conn.commit()
    conn.close()
    st.success("🗑️ Deleted successfully!")
    st.experimental_rerun()

# ------------------- Admin Dashboard -------------------
def admin_dashboard():
    st.markdown("## 👑 Admin Dashboard")
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()

    c.execute("SELECT COUNT(*) FROM users")
    total_users = c.fetchone()[0]

    c.execute("SELECT COUNT(*) FROM feedback")
    total_feedbacks = c.fetchone()[0]

    c.execute("SELECT role, COUNT(*) FROM feedback GROUP BY role")
    role_data = c.fetchall()

    conn.close()

    st.metric("👥 Total Users", total_users)
    st.metric("📄 Total Feedbacks", total_feedbacks)

    roles = [r[0] for r in role_data]
    counts = [r[1] for r in role_data]
    st.bar_chart(data={"Roles": roles, "Feedbacks": counts})


#8
# ------------------- Change Password -------------------
def change_password():
    st.markdown("## 🔐 Change Password")
    with st.form("change_password_form"):
        current_pw = st.text_input("Current Password", type="password", key="curr_pw_change")
        new_pw = st.text_input("New Password", type="password", key="new_pw_change")
        confirm_pw = st.text_input("Confirm New Password", type="password", key="confirm_pw_change")
        submitted = st.form_submit_button("🔁 Update Password")
        if submitted:
            if new_pw != confirm_pw:
                st.error("❌ New passwords do not match.")
                return
            conn = sqlite3.connect(DB_FILE)
            c = conn.cursor()
            c.execute("SELECT * FROM users WHERE email=? AND password=?", (st.session_state.email, current_pw))
            result = c.fetchone()
            if result:
                c.execute("UPDATE users SET password=? WHERE email=?", (new_pw, st.session_state.email))
                conn.commit()
                st.success("✅ Password updated successfully!")
            else:
                st.error("❌ Incorrect current password.")
            conn.close()

# ------------------- Forgot Password -------------------
def forgot_password():
    st.markdown("## ❓ Forgot Password")
    with st.form("forgot_pw_form"):
        email = st.text_input("Enter your registered Email")
        method = st.radio("Verify using:", ["📧 Email", "📱 Phone"])
        submitted = st.form_submit_button("Send OTP")
        if submitted:
            otp = generate_otp()
            st.session_state.reset_email = email
            st.session_state.generated_otp = otp
            st.session_state.otp_verified = False
            if method == "📧 Email":
                send_email_otp(email, otp)
            else:
                st.warning("📱 Phone OTP not implemented. Placeholder only.")
            st.info("📨 OTP sent. Please check and enter below.")

    if "generated_otp" in st.session_state:
        user_otp = st.text_input("Enter OTP", max_chars=6)
        if st.button("✅ Verify OTP"):
            if user_otp == st.session_state.generated_otp:
                st.success("✅ OTP Verified. You may now reset your password.")
                st.session_state.otp_verified = True
            else:
                st.error("❌ Invalid OTP.")

    if st.session_state.get("otp_verified"):
        with st.form("reset_pw_form"):
            new_pw = st.text_input("New Password", type="password", key="reset_pw1")
            confirm_pw = st.text_input("Confirm New Password", type="password", key="reset_pw2")
            reset = st.form_submit_button("🔄 Reset Password")
            if reset:
                if new_pw == confirm_pw:
                    conn = sqlite3.connect(DB_FILE)
                    c = conn.cursor()
                    c.execute("UPDATE users SET password=? WHERE email=?", (new_pw, st.session_state.reset_email))
                    conn.commit()
                    conn.close()
                    st.success("✅ Password reset successful. Please login again.")
                    st.session_state.page = "login"
                else:
                    st.error("❌ Passwords do not match.")

# ------------- OTP Email Placeholder -------------
def send_email_otp(email, otp):
    try:
        msg = EmailMessage()
        msg.set_content(f"Your OTP for Resume Feedback App is: {otp}")
        msg["Subject"] = "OTP Verification - Resume Feedback App"
        msg["From"] = EMAIL_ADDRESS
        msg["To"] = email
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        server.send_message(msg)
        server.quit()
        st.success("✅ OTP sent to email.")
    except Exception as e:
        st.error(f"❌ Failed to send email: {e}")


#9
# ---------------- Dashboard Navigation & Routing ----------------
def run_dashboard():
    st.sidebar.image("https://cdn-icons-png.flaticon.com/512/1077/1077063.png", width=100)
    st.sidebar.markdown(f"### 👋 Welcome, **{st.session_state.name}**")
    dashboard_menu = st.sidebar.radio("📋 Menu", [
        "📤 Upload Resume",
        "📊 Resume Feedback",
        "🔄 Rewritten Resume",
        "🔈 Audio Tips",
        "📧 Email Feedback",
        "📂 My History",
        "🔐 Change Password",
        "👑 Admin Dashboard",
        "❓ Forgot Password",
        "🚪 Logout"
    ])

    if dashboard_menu == "📤 Upload Resume":
        upload_resume()

    elif dashboard_menu == "📊 Resume Feedback":
        display_feedback()

    elif dashboard_menu == "🔄 Rewritten Resume":
        rewrite_resume()

    elif dashboard_menu == "🔈 Audio Tips":
        generate_audio_tips()

    elif dashboard_menu == "📧 Email Feedback":
        send_feedback_email()

    elif dashboard_menu == "📂 My History":
        show_history()

    elif dashboard_menu == "🔐 Change Password":
        change_password()

    elif dashboard_menu == "👑 Admin Dashboard":
        admin_dashboard()

    elif dashboard_menu == "❓ Forgot Password":
        forgot_password()

    elif dashboard_menu == "🚪 Logout":
        st.success("✅ Logged out successfully.")
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.experimental_rerun()

    st.markdown("---")
    st.markdown("##### 💼 Built with ❤️ by SHATU | Powered by GPT, Groq & Streamlit")
    st.markdown('<style>.block-container{padding-top:2rem;}</style>', unsafe_allow_html=True)

# ---------------- Start App ----------------
if __name__ == "__main__":
    init_db()
    if not st.session_state.get("logged_in"):
        run_login_signup()
    else:
        run_dashboard()
